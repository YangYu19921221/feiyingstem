#!/usr/bin/env python3
"""把《音标教学流程-老师审阅版.md》渲染成 docx(老师微信转发用)。

一次性脚本:只处理该文档用到的 Markdown 子集(#标题/表格/列表/引用/代码块/粗体)。
用法: python3 docs/gen_phonetic_teacher_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SRC = Path(__file__).parent / "音标教学流程-老师审阅版.md"
DST = Path(__file__).parent / "音标教学流程-老师审阅版.docx"

ACCENT = RGBColor(0xBD, 0x52, 0x27)  # 品牌暖橙


def set_cn_font(run, size=None, bold=None, color=None, mono=False):
    font = "Courier New" if mono else "微软雅黑"
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_rich(paragraph, text, size=10.5, mono=False):
    """**粗体** 拆段渲染,其余原样。"""
    for i, seg in enumerate(re.split(r"\*\*", text)):
        if not seg:
            continue
        run = paragraph.add_run(seg)
        set_cn_font(run, size=size, bold=(i % 2 == 1), mono=mono)


def main():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    # 页面默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)

    i, in_code = 0, False
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            p = doc.add_paragraph()
            add_rich(p, line or " ", size=9, mono=True)
            i += 1
            continue

        # 表格:收集连续的 | 行
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", c) for c in cells):  # 跳过分隔行
                    rows.append(cells)
                i += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 2"
                for r, row in enumerate(rows):
                    for c, cell_text in enumerate(row):
                        cell = table.cell(r, min(c, len(table.columns) - 1))
                        cell.paragraphs[0].text = ""
                        add_rich(cell.paragraphs[0], cell_text, size=9)
                        if r == 0:
                            for run in cell.paragraphs[0].runs:
                                run.font.bold = True
                doc.add_paragraph()
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            set_cn_font(run, size=18, bold=True, color=ACCENT)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            set_cn_font(run, size=14, bold=True, color=ACCENT)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            set_cn_font(run, size=12, bold=True)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            add_rich(p, line[2:].strip(), size=10)
            for run in p.runs:
                run.font.italic = True
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_rich(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich(p, line[2:])
        elif line.strip() in ("---", "*", ""):
            pass  # 分隔线/空行:docx 里靠段距,不加空段
        else:
            p = doc.add_paragraph()
            add_rich(p, re.sub(r"^\*|\*$", "", line.strip()))
        i += 1

    doc.save(DST)
    print(f"OK -> {DST}")


if __name__ == "__main__":
    main()
