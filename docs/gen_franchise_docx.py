# -*- coding: utf-8 -*-
"""生成加盟方案 Word 文档(对外招商版 + 对内成本版)。"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 品牌配色
ORANGE = RGBColor(0xFF, 0x6B, 0x35)   # 活力橙
YELLOW = RGBColor(0xF5, 0xA0, 0x1F)   # 暖黄(比 FFD23F 深,保证白底可读)
BLUE   = RGBColor(0x00, 0x9E, 0xC4)   # 天空蓝(加深)
DARK   = RGBColor(0x2B, 0x2B, 0x2B)   # 正文深灰
GRAY   = RGBColor(0x88, 0x88, 0x88)   # 次要灰
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
CN_FONT = "微软雅黑"


def set_cn_font(run, name=CN_FONT):
    run.font.name = name
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), name)


def shade_cell(cell, hex_color):
    """给单元格加底色。hex_color 如 'FF6B35'。"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=DARK, size=10.5,
                  align=WD_ALIGN_PARAGRAPH.LEFT, font=CN_FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    set_cn_font(run, font)


def add_title(doc, text, color=ORANGE, size=15):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    set_cn_font(run)
    # 底部橙线
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "FF6B35")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def add_sub(doc, text, color=BLUE, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    set_cn_font(run)
    return p


def add_body(doc, text, *, size=10.5, color=DARK, bold=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    set_cn_font(run)
    return p


def add_bullet(doc, text, *, color=DARK, size=10.5):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    r0 = p.add_run("• ")
    r0.font.color.rgb = ORANGE
    r0.bold = True
    r0.font.size = Pt(size)
    set_cn_font(r0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    set_cn_font(run)
    return p


def add_table(doc, headers, rows, *, header_fill="FF6B35", col_widths=None,
              zebra="FFF3EC"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    # 表头
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_text(c, h, bold=True, color=WHITE, size=10.5,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(c, header_fill)
    # 数据行
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[ci], str(val), size=10, align=align)
            if ri % 2 == 1 and zebra:
                shade_cell(cells[ci], zebra)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t


def add_callout(doc, text, fill="FFF3EC", bar="FF6B35"):
    """引用/强调框:单格表格加左侧色条效果(用底色近似)。"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    shade_cell(c, fill)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    run.bold = True
    set_cn_font(run)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def base_doc():
    doc = Document()
    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = CN_FONT
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    # 页边距
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
    return doc


def add_cover(doc, main_title, subtitle, tag):
    # 顶部留白
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("🎯")
    r.font.size = Pt(48)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(main_title)
    r2.bold = True
    r2.font.size = Pt(30)
    r2.font.color.rgb = ORANGE
    set_cn_font(r2)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    r3.bold = True
    r3.font.size = Pt(16)
    r3.font.color.rgb = BLUE
    set_cn_font(r3)
    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(tag)
    r4.font.size = Pt(12)
    r4.font.color.rgb = GRAY
    set_cn_font(r4)
    for _ in range(6):
        doc.add_paragraph()
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("云南首批合作 · 起步期特惠")
    rf.font.size = Pt(12)
    rf.bold = True
    rf.font.color.rgb = YELLOW
    set_cn_font(rf)
    doc.add_page_break()


# ============ 文档一:对外招商版 ============
def build_external(path):
    doc = base_doc()
    add_cover(doc, "智能英语学习系统", "加盟合作方案",
              "教 · 学 · 测 · 评 · 赛  一体化智能教学系统")

    add_title(doc, "一、我们是什么")
    add_body(doc, "我们不是又一个“背单词 APP”,而是一套 “教-学-测-评-赛” 完整闭环的智能英语教学系统。")
    add_body(doc, "中小机构不用自己开发、不用请技术团队,开箱即拥有大机构级别的数字化教学能力和游戏化留生工具:")
    add_bullet(doc, "老师端:班级管理、内容自建、精细化分配、作业布置、实时课堂监控、投屏大屏")
    add_bullet(doc, "学生端:闪卡 / 听写 / 拼写 / 选择 / 填空 / 例句 / 句子背诵 / 单元考试 等多种模式")
    add_bullet(doc, "游戏化:宠物养成、真人实时 PK 对战、全自动晋级赛、段位 / 成就 / 排行榜")
    add_bullet(doc, "AI 能力:自动出题、智能组卷、千人千卷、错因讲解、学习周报")
    add_bullet(doc, "发音:剑桥真人发音 + 双引擎语音评测,不请外教也能练口语")
    add_bullet(doc, "家长端:绑定孩子看学习看板 + 周报;招生工具:AI 英语测评漏斗,自带获客能力")
    add_callout(doc, "别人加盟给你一套课件,我们给你一套能自己跑起来的数字化教学系统。")

    add_title(doc, "二、核心卖点")
    add_body(doc, "机构老板最关心三件事:招得来、留得住、教得省。系统正好对着这三点。", bold=True)
    add_sub(doc, "① 帮你招生(获客)")
    add_bullet(doc, "AI 英语水平测评漏斗:一次匿名测评 → 专业深度报告 → 自动进招生线索池,老师跟进转化")
    add_bullet(doc, "地推、朋友圈、进校、家长会都能用,自带免费获客入口")
    add_sub(doc, "② 帮你留生(续费)")
    add_bullet(doc, "宠物养成 + 真人实时 PK + 全自动晋级赛 + 段位 / 成就 / 排行榜,黏性远超普通打卡 APP")
    add_bullet(doc, "防划水 + 切屏专注监控 + 家长端周报,家长看得见真实学习,续费有据可讲")
    add_sub(doc, "③ 帮你省人省课时(降本)")
    add_bullet(doc, "AI 自动出题、智能组卷、千人千卷、错因讲解、周报,备课批改全包")
    add_bullet(doc, "真人级发音评测,不请外教也能练口语跟读")
    add_bullet(doc, "一个老师能带的学生数直接翻倍")

    doc.add_page_break()
    add_title(doc, "三、加盟费用结构")
    add_callout(doc, "定价逻辑:加盟门槛明显低于同行,让本地机构零压力进场;赠送账号控成本,超出部分从第一年起按人头收费——你招得越多我们才多赚。")

    add_sub(doc, "线一:单校区加盟(主力 · 4 万区间)")
    add_table(doc,
              ["档位", "加盟费", "含首年账号", "适合"],
              [["启航档", "¥29,800", "含 60 名", "小校区试水"],
               ["标准档(主推)", "¥39,800", "含 100 名", "常规单校区机构"]],
              col_widths=[4.0, 3.0, 3.5, 5.0])
    add_body(doc, "单校区档保护范围为“本校区自用”,不排他,主打低门槛铺量。", size=9.5, color=GRAY)

    add_sub(doc, "线二:区域独家代理(走利润 · 明显更贵)")
    add_body(doc, "核心区别:在授权区域内独家,公司不再发展第二家;代理可在辖区内发展下级校区、参与分成。")
    add_table(doc,
              ["档位", "独家范围", "加盟费", "含首年账号", "权益"],
              [["县级独家", "1 个县/区", "¥98,000", "300 名", "辖区独家 + 招下级抽成"],
               ["市级独家", "1 个地级市", "¥198,000", "600 名", "全市独家 + 下级分成 + 优先新功能"],
               ["州/省级", "自治州/全省", "面议(30万+)", "面议", "最高级别,一事一议"]],
              col_widths=[3.0, 3.2, 3.2, 2.8, 4.8])
    add_body(doc, "同行区域独家常见 30-80 万起,本方案县级 9.8 万、市级 19.8 万,仍明显低于行情。",
             size=9.5, color=GRAY)

    add_sub(doc, "加盟费一口价全含(优惠核心)")
    for x in ["品牌授权 + 系统开通", "首年赠送学生账号(见上表)", "老师全员培训 + 初始词库 / 班级代建",
              "首年 AI 额度(合理上限内)全包", "全套招生物料 + AI 测评获客工具 + 话术",
              "系统终身免费更新升级 + 专属客服群"]:
        add_bullet(doc, "✅ " + x)

    add_sub(doc, "超出赠送账号 → 从第一年起按人头收费(元 / 生 / 年)")
    add_table(doc,
              ["学生规模(超出部分)", "单价"],
              [["第 101 - 300 人", "60"],
               ["第 301 - 500 人", "50"],
               ["第 501 人以上", "40"]],
              col_widths=[8.0, 4.0])
    add_body(doc, "折算一个学生一年几十块、一个月几块钱,机构一节课时费就覆盖了。招得越多单价越低。",
             size=9.5, color=GRAY)

    add_sub(doc, "续费(次年起)")
    add_table(doc,
              ["档位", "年服务费", "超出账号"],
              [["单校区标准档", "¥6,800 / 年", "同上阶梯价"],
               ["县级独家", "¥12,800 / 年", "同上阶梯价"],
               ["市级独家", "¥19,800 / 年", "同上阶梯价"]],
              col_widths=[4.5, 4.0, 4.0])

    add_sub(doc, "算笔账(单校区标准档)")
    add_table(doc,
              ["场景", "首年", "次年起"],
              [["招到 100 人(用满赠送)", "39,800", "6,800"],
               ["招到 200 人", "45,800", "12,800"],
               ["招到 300 人", "51,800", "18,800"]],
              col_widths=[6.0, 3.5, 3.5])
    add_callout(doc, "话术:3 万 9 含 100 个学生账号全年全含,超出每人一年才几十块——你招得越多我们越便宜,你没做大我们也不多收一分。")

    doc.add_page_break()
    add_title(doc, "四、我们送什么")
    add_sub(doc, "送服务")
    for x in ["老师全员线上培训 + 考核(教会为止)", "系统开通、初始班级 / 词库导入代建",
              "首年 AI 额度(合理上限内)全包", "系统持续免费更新升级,新功能自动到账",
              "专属客服群 + 技术响应", "招生旺季运营方法培训"]:
        add_bullet(doc, "✅ " + x)
    add_sub(doc, "送物料")
    for x in ["招生海报 / 朋友圈图 / 家长会 PPT 模板(可印自有 logo)",
              "AI 测评获客活动全套物料 + 话术", "家长沟通话术、续费话术", "课堂大屏投屏方案(现成)"]:
        add_bullet(doc, "✅ " + x)

    add_title(doc, "五、AI 使用说明与额度")
    add_sub(doc, "A 类:一次生成、全校复用(不计入用量)")
    add_body(doc, "AI 例句、干扰项、音标释义补全、单词本封面——同一单词全校共享一份,不设上限。")
    add_sub(doc, "B 类:按人头实时消耗(设合理上限,正常教学完全够用)")
    add_table(doc,
              ["功能", "触发场景", "额度上限"],
              [["AI 错因讲解", "学生答错时讲解", "每生每日 20 次"],
               ["AI 智能组卷", "老师一键出卷", "每位老师每日 10 份"],
               ["AI 竞赛出题", "批量生成竞赛题", "每位老师每月 200 题"],
               ["AI 学习周报", "家长/老师查看", "每生每周 1 份(自动缓存)"],
               ["AI 招生测评报告", "潜在学员测评", "每加盟点每月 500 份"]],
              col_widths=[3.5, 5.0, 4.0])
    add_body(doc, "宠物对战、PK 竞技场的“AI 对手”为智能算法生成,不消耗大模型额度,可无限畅玩。",
             size=9.5, color=GRAY)

    add_title(doc, "六、加盟商需具备 / 投入")
    for x in ["有实际办学场地或线上招生能力", "每个班配一块屏 / 一台可投屏设备(课堂大屏用)",
              "学生端设备:平板或手机(家长自备即可)", "遵守品牌规范、不跨区窜货(区域独家档受保护)"]:
        add_bullet(doc, x)

    add_title(doc, "七、合作流程")
    for i, x in enumerate(["咨询沟通,确认档位与区域", "签订加盟协议、缴纳加盟费",
                           "系统开通 + 老师培训 + 词库代建", "领取招生物料,启动首波招生",
                           "持续运营支持,旺季运营培训"], 1):
        add_bullet(doc, f"{i}. {x}")

    doc.add_paragraph()
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run("— 期待与你在云南共建智能英语教育品牌 —")
    rf.font.color.rgb = ORANGE
    rf.bold = True
    set_cn_font(rf)

    doc.save(path)
    print("已生成:", path)


# ============ 文档二:对内成本版 ============
def build_internal(path):
    doc = base_doc()
    p = doc.add_paragraph()
    r = p.add_run("【内部机密 · 不对外发放】")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)
    set_cn_font(r)
    add_title(doc, "加盟方案 · 对内成本与毛利测算", color=RGBColor(0xC0, 0x30, 0x30))

    add_sub(doc, "单个学生一年的真实 AI 成本")
    add_body(doc, "成本几乎完全取决于后台配置的模型(存于数据库 ai_providers 表,需登录管理后台查看)。")
    add_table(doc,
              ["后台配置模型", "单生年 AI 成本", "100 赠送账号年硬成本", "备注"],
              [["通义千问 / 国产模型", "几元以内", "< ¥1,000", "代码 DashScope 配置暗示线上很可能用此类"],
               ["Claude Sonnet", "约 ¥30", "约 ¥3,000", "—"],
               ["GPT-4-turbo", "约 ¥60-70", "约 ¥6,000-7,000", "个性化组卷 16000 token,成本最高"]],
              header_fill="C03030", zebra="F7E9E9",
              col_widths=[4.0, 3.5, 4.0, 5.0])
    add_callout(doc, "关键动作:登录管理后台确认 ai_providers 默认启用的模型。若为国产模型,送 100 账号成本可忽略,标准档稳赚。",
                fill="FBEDED")

    add_sub(doc, "标准档毛利粗算(假设线上用国产模型)")
    add_table(doc,
              ["项目", "首年", "次年起"],
              [["收入(招 100 人)", "39,800", "6,800"],
               ["AI 硬成本", "< 1,000", "< 1,000"],
               ["服务器 / 带宽(分摊)", "视规模", "视规模"],
               ["培训 / 物料(一次性)", "人力成本", "—"],
               ["毛利", "≈ 3.5 万+", "≈ 5,000+"]],
              header_fill="C03030", zebra="F7E9E9",
              col_widths=[6.0, 3.5, 3.5])

    add_sub(doc, "三个必做事项(定价才踏实)")
    add_body(doc, "1. 确认线上模型", bold=True)
    add_bullet(doc, "后台查 ai_providers 默认模型 → 决定送账号是否亏")
    add_body(doc, "2. 补齐 AI 限流(当前代码对 AI 零配额控制,任何用户可无限触发)", bold=True)
    add_bullet(doc, "explain_mistake(错因讲解):学生端高频、零缓存 → 加每生每日上限 + 持久缓存")
    add_bullet(doc, "generate_personalized_exam(个性化组卷):单次 16000 token → 加每老师每日上限")
    add_bullet(doc, "generate_competition_question(竞赛出题):循环调用 → 加每月总量上限")
    add_body(doc, "3. 修缓存", bold=True)
    add_bullet(doc, "当前是进程内存缓存(重启失效、多进程不共享),ai_cache 表未被使用")
    add_bullet(doc, "建议改 Redis 或落库,让 A 类内容真正“一次生成全校复用”,否则实际成本高于上表估算")

    add_sub(doc, "涨价路线图(内部)")
    add_bullet(doc, "第一阶段(首批 30 家):引流价,主打云南本地口碑案例")
    add_bullet(doc, "第二阶段(品牌立住):加盟费上浮,赠送账号缩减,人头费单价上调")
    add_bullet(doc, "第三阶段(跨省):恢复行业正常加盟费;老加盟商锁定老价续费,制造“早进早赚”")

    add_sub(doc, "合同关键条款清单(签约前写死)")
    for x in ["授权范围:单校区自用 / 区域独家边界(精确到县/市)",
              "独家期限:建议 1-2 年一签,约定“做不到量可收回独家权”",
              "账号性质:赠送账号为“首年”还是“永久”,次年如何续、超出如何计费",
              "下级分成(区域档):抽成比例与结算方式",
              "数据归属:学生数据归属、隐私合规、退出时数据处理",
              "违约与退出:退款规则、终止条件",
              "价格调整权:注明首批价的有效范围"]:
        add_bullet(doc, x)

    doc.save(path)
    print("已生成:", path)


if __name__ == "__main__":
    import os
    here = os.path.dirname(os.path.abspath(__file__))
    build_external(os.path.join(here, "加盟合作方案(对外招商版).docx"))
    build_internal(os.path.join(here, "加盟成本测算(对内机密版).docx"))
